import pandas as pd
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import logging
import os
import json
from typing import Optional
from dotenv import load_dotenv
import openpyxl

load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class DocumentParser:
    """
    A service for extracting and parsing content from various document types.
    It's designed to be a standalone utility class that does not hold references to other services.
    """

    def extract_text(self, file_path: str) -> str:
        """
        Extracts raw text content from PDF, DOCX, TXT, and CSV files, including tables.
        
        Args:
            file_path: The path to the document file.

        Returns:
            The extracted text as a single string.
        """
        _, file_extension = os.path.splitext(file_path.lower())
        text = ""

        if file_extension == '.pdf':
            try:
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            except Exception as e:
                logging.error(f"Failed to extract text from PDF file: {e}")
        
        elif file_extension == '.docx':
            try:
                doc = DocxDocument(file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " | "
                        text += "\n"
            except Exception as e:
                logging.error(f"Failed to extract text from DOCX file: {e}")
        
        elif file_extension == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except Exception as e:
                logging.error(f"Failed to read TXT file: {e}")

        elif file_extension == '.csv':
            try:
                # Read as a DataFrame and convert to string for consistent LLM input
                df = pd.read_csv(file_path, encoding='utf-8')
                text = df.to_string()
            except Exception as e:
                logging.error(f"Failed to read CSV file: {e}")
        
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return text

    def parse_item_master(self, file_path: str, gemini_service) -> list:
        """
        Parses item master content from a CSV or Excel file and standardizes
        column names using an LLM API call. The gemini_service is passed here
        at call time, not during initialization, to avoid a circular dependency.
        
        Args:
            file_path: The path to the item master file.
            gemini_service: The GeminiAgentService instance.
            
        Returns:
            A list of dictionaries with standardized column names.
        """
        try:
            _, file_extension = os.path.splitext(file_path.lower())
            
            if file_extension == '.csv':
                df = pd.read_csv(file_path, encoding='utf-8')
            elif file_extension in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path, engine='openpyxl')
            else:
                raise ValueError(f"Unsupported file type for item master: {file_extension}")
            
            # Use LLM to map columns to standard ones
            csv_content = df.to_csv(index=False)
            logging.info("Calling LLM to standardize item master columns...")
            standardized_data = gemini_service.standardize_item_master(csv_content)
            
            return standardized_data
        except Exception as e:
            logging.error(f"Error parsing item master file with LLM: {e}")
            return []